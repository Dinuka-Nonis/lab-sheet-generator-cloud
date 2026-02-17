"""
SLIIT Template - Professional university format
"""

from templates.base import BaseTemplate
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class SLIITTemplate(BaseTemplate):
    """SLIIT official template with university branding."""
    
    def __init__(self):
        super().__init__()
        self.template_name = "SLIIT"
        self.template_id = "sliit"
    
    def generate(self, student_name, student_id, module_name, module_code, 
                 sheet_label, logo_path=None):
        """Generate document using SLIIT template."""
        
        # Create document
        doc = self._create_document()
        
        # Set narrow margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Header with logo and university name
        if logo_path:
            self._add_logo(doc, logo_path, width=1.5)
        
        # University name
        uni_name = doc.add_paragraph()
        uni_run = uni_name.add_run("SRI LANKA INSTITUTE OF INFORMATION TECHNOLOGY")
        uni_run.font.size = Pt(16)
        uni_run.font.bold = True
        uni_run.font.color.rgb = RGBColor(0, 51, 102)  # SLIIT blue
        uni_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Faculty
        faculty = doc.add_paragraph()
        faculty_run = faculty.add_run("Faculty of Computing")
        faculty_run.font.size = Pt(12)
        faculty_run.font.color.rgb = RGBColor(0, 51, 102)
        faculty.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacing
        
        # Horizontal line
        self._add_horizontal_line(doc)
        
        doc.add_paragraph()  # Spacing
        
        # Module name - Large and bold
        module_para = doc.add_paragraph()
        module_run = module_para.add_run(module_name)
        module_run.font.size = Pt(18)
        module_run.font.bold = True
        module_run.font.color.rgb = RGBColor(0, 51, 102)
        module_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Module code
        code_para = doc.add_paragraph()
        code_run = code_para.add_run(f"({module_code})")
        code_run.font.size = Pt(14)
        code_run.font.color.rgb = RGBColor(102, 102, 102)
        code_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacing
        
        # Sheet label - Highlighted
        label_para = doc.add_paragraph()
        label_run = label_para.add_run(sheet_label)
        label_run.font.size = Pt(16)
        label_run.font.bold = True
        label_run.font.color.rgb = RGBColor(204, 0, 0)  # Red highlight
        label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacing
        
        # Horizontal line
        self._add_horizontal_line(doc)
        
        doc.add_paragraph()  # Spacing
        
        # Student information table - Professional style
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Set column widths
        for row in table.rows:
            row.cells[0].width = Inches(2)
            row.cells[1].width = Inches(4)
        
        # Student info
        info = [
            ('Student Name:', student_name),
            ('Student ID:', student_id),
            ('Date:', self._get_current_date())
        ]
        
        for i, (label, value) in enumerate(info):
            # Label cell
            label_cell = table.rows[i].cells[0]
            label_para = label_cell.paragraphs[0]
            label_para.text = label
            label_run = label_para.runs[0]
            label_run.font.bold = True
            label_run.font.size = Pt(11)
            label_run.font.color.rgb = RGBColor(0, 51, 102)
            
            # Value cell
            value_cell = table.rows[i].cells[1]
            value_para = value_cell.paragraphs[0]
            value_para.text = value
            value_run = value_para.runs[0]
            value_run.font.size = Pt(11)
            
            # Shading for header
            if i == 0:
                self._set_cell_background(label_cell, "E7F0FF")
                self._set_cell_background(value_cell, "F5F9FF")
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Content sections with SLIIT styling
        sections = [
            {
                'title': 'Learning Objectives',
                'content': 'State the learning objectives for this practical session.'
            },
            {
                'title': 'Tasks',
                'content': '1. Describe the first task to be completed.\n2. Describe the second task.\n3. Describe the third task.\n4. Add more tasks as needed.'
            },
            {
                'title': 'Procedure',
                'content': 'Document the step-by-step procedure you followed to complete the tasks.'
            },
            {
                'title': 'Results and Observations',
                'content': 'Present your results, findings, and observations here. Include screenshots, diagrams, or tables as appropriate.'
            },
            {
                'title': 'Discussion',
                'content': 'Discuss your results. Explain what you learned and any challenges you encountered.'
            },
            {
                'title': 'Conclusion',
                'content': 'Summarize what you learned from this practical session.'
            }
        ]
        
        for section in sections:
            # Section heading with SLIIT blue
            heading = doc.add_heading(level=2)
            heading_run = heading.add_run(section['title'])
            heading_run.font.color.rgb = RGBColor(0, 51, 102)
            heading_run.font.size = Pt(14)
            heading_run.font.bold = True
            
            # Content
            content_para = doc.add_paragraph(section['content'])
            content_para.style = 'Normal'
            
            # Add spacing
            doc.add_paragraph()
        
        # Footer
        doc.add_paragraph()
        self._add_horizontal_line(doc)
        
        footer = doc.add_paragraph()
        footer_run = footer.add_run("© Sri Lanka Institute of Information Technology")
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Generate filename
        filename = f"{sheet_label.replace(' ', '_')}_{student_id}.docx"
        
        # Save
        return self._save_document(doc, filename)
    
    def _add_horizontal_line(self, doc):
        """Add a horizontal line."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add bottom border to paragraph
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        pPr.insert_element_before(pBdr, 'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
                                   'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
                                   'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
                                   'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
                                   'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap', 'w:outlineLvl',
                                   'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr', 'w:pPrChange')
        
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '003366')
        pBdr.append(bottom)
    
    def _set_cell_background(self, cell, color):
        """Set cell background color."""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color)
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    def _get_current_date(self):
        """Get current date formatted."""
        from datetime import datetime
        return datetime.now().strftime('%B %d, %Y')
