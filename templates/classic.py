"""
Classic template - Simple and clean design
"""

from templates.base import BaseTemplate
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


class ClassicTemplate(BaseTemplate):
    """Classic template with clean, professional layout."""
    
    def __init__(self):
        super().__init__()
        self.template_name = "Classic"
        self.template_id = "classic"
    
    def generate(self, student_name, student_id, module_name, module_code, 
                 sheet_label, logo_path=None):
        """Generate document using classic template."""
        
        # Create document
        doc = self._create_document()
        
        # Add logo if provided
        self._add_logo(doc, logo_path, width=1.2)
        
        # Add spacing
        doc.add_paragraph()
        
        # Title
        self._add_title(doc, f"{module_name}", size=18, bold=True)
        self._add_title(doc, sheet_label, size=16, bold=True)
        
        # Add spacing
        doc.add_paragraph()
        
        # Student information table
        self._add_info_table(doc, student_name, student_id, module_name, 
                            module_code, sheet_label)
        
        # Add spacing
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Content sections
        sections = [
            {
                'title': 'Objective',
                'content': 'Write the objective of this practical session here.'
            },
            {
                'title': 'Tasks',
                'content': '1. Task description here\n2. Another task\n3. Final task'
            },
            {
                'title': 'Procedure',
                'content': 'Describe the procedure followed during the practical.'
            },
            {
                'title': 'Results',
                'content': 'Document your results and observations here.'
            },
            {
                'title': 'Conclusion',
                'content': 'Write your conclusion here.'
            }
        ]
        
        self._add_sections(doc, sections)
        
        # Generate filename
        filename = f"{sheet_label.replace(' ', '_')}_{student_id}.docx"
        
        # Save
        return self._save_document(doc, filename)
