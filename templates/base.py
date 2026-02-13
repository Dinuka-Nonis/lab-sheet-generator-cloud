"""
Base template class for document generation
Adapted for cloud service
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from abc import ABC, abstractmethod
import os
from datetime import datetime


class BaseTemplate(ABC):
    """Base class for all document templates."""
    
    def __init__(self):
        self.template_name = "Base Template"
        self.template_id = "base"
    
    @abstractmethod
    def generate(self, student_name, student_id, module_name, module_code, 
                 sheet_label, logo_path=None):
        """Generate document. Must be implemented by subclasses."""
        pass
    
    def _create_document(self):
        """Create new document."""
        return Document()
    
    def _add_logo(self, doc, logo_path, width=1.2):
        """Add logo to document."""
        if logo_path and os.path.exists(logo_path):
            try:
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(logo_path, width=Inches(width))
                return True
            except Exception as e:
                print(f"Warning: Could not add logo: {e}")
                return False
        return False
    
    def _add_title(self, doc, text, size=16, bold=True, align='center'):
        """Add title paragraph."""
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
        
        if align == 'center':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == 'left':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif align == 'right':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        return paragraph
    
    def _add_info_table(self, doc, student_name, student_id, module_name, 
                        module_code, sheet_label):
        """Add information table."""
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Light List Accent 1'
        
        # Headers and values
        info = [
            ('Student Name', student_name),
            ('Student ID', student_id),
            ('Module', f"{module_name} ({module_code})"),
            ('Sheet Type', sheet_label),
            ('Date', datetime.now().strftime('%Y-%m-%d'))
        ]
        
        for i, (label, value) in enumerate(info):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value
            
            # Bold labels
            table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        
        return table
    
    def _add_sections(self, doc, sections):
        """Add content sections."""
        for section in sections:
            # Section heading
            heading = doc.add_heading(section['title'], level=2)
            
            # Section content
            if 'content' in section:
                doc.add_paragraph(section['content'])
            
            # Add space
            doc.add_paragraph()
    
    def _save_document(self, doc, output_filename):
        """Save document to file."""
        doc.save(output_filename)
        return output_filename
